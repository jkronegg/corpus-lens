import importlib.util
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SCRIPT_PATH = (
    Path(__file__).resolve().parents[1]
    / "scripts"
    / "convert_md_to_docx.py"
)

spec = importlib.util.spec_from_file_location("convert_md_to_docx", SCRIPT_PATH)
module = importlib.util.module_from_spec(spec)
assert spec and spec.loader
spec.loader.exec_module(module)


class ConvertMarkdownToDocxTests(unittest.TestCase):
    def convert(self, markdown: str) -> Document:
        with tempfile.TemporaryDirectory() as tmpdir:
            md_path = Path(tmpdir) / "input.md"
            docx_path = Path(tmpdir) / "output.docx"
            md_path.write_text(markdown, encoding="utf-8")
            module.parse_markdown_to_docx(str(md_path), str(docx_path))
            return Document(docx_path)

    def test_headings_are_mapped_to_word_heading_styles(self) -> None:
        doc = self.convert("# Titre 1\n## Titre 2\n### Titre 3\n")

        self.assertEqual(doc.paragraphs[0].text, "Titre 1")
        self.assertEqual(doc.paragraphs[0].style.name, "Heading 1")
        self.assertEqual(doc.paragraphs[1].text, "Titre 2")
        self.assertEqual(doc.paragraphs[1].style.name, "Heading 2")
        self.assertEqual(doc.paragraphs[2].text, "Titre 3")
        self.assertEqual(doc.paragraphs[2].style.name, "Heading 3")

    def test_bullet_list_items_use_list_bullet_style_and_indent(self) -> None:
        doc = self.convert("- Premier\n  - Enfant\n* Deuxieme\n")

        self.assertEqual(doc.paragraphs[0].style.name, "List Bullet")
        self.assertEqual(doc.paragraphs[0].text, "Premier")
        self.assertEqual(doc.paragraphs[0].paragraph_format.left_indent, 0)

        self.assertEqual(doc.paragraphs[1].style.name, "List Bullet")
        self.assertEqual(doc.paragraphs[1].text, "Enfant")
        nested_indent = doc.paragraphs[1].paragraph_format.left_indent
        self.assertIsNotNone(nested_indent)
        self.assertGreater(nested_indent, 700000)
        self.assertLess(nested_indent, 730000)

        self.assertEqual(doc.paragraphs[2].style.name, "List Bullet")
        self.assertEqual(doc.paragraphs[2].text, "Deuxieme")

    def test_numbered_list_items_use_list_number_style(self) -> None:
        doc = self.convert("1) Un\n2) Deux\n")

        self.assertEqual(doc.paragraphs[0].style.name, "List Number")
        self.assertEqual(doc.paragraphs[0].text, "Un")
        self.assertEqual(doc.paragraphs[1].style.name, "List Number")
        self.assertEqual(doc.paragraphs[1].text, "Deux")

    def test_inline_link_creates_hyperlink_relationship(self) -> None:
        url = "https://example.org"
        doc = self.convert(f"Lien: [Example]({url})\n")

        paragraph = doc.paragraphs[0]
        hyperlinks = paragraph._p.findall(qn("w:hyperlink"))
        self.assertEqual(len(hyperlinks), 1)

        rel_id = hyperlinks[0].get(qn("r:id"))
        self.assertIsNotNone(rel_id)
        relationship = paragraph.part.rels[rel_id]
        self.assertEqual(relationship.target_ref, url)

    def test_markdown_table_is_kept_as_text_lines(self) -> None:
        markdown = "| Col1 | Col2 |\n| --- | --- |\n| A | B |\n"
        doc = self.convert(markdown)

        self.assertEqual(len(doc.tables), 0)
        self.assertEqual(doc.paragraphs[0].text, "| Col1 | Col2 | | --- | --- | | A | B |")

    def test_inline_bold_italic_and_code_runs_are_styled(self) -> None:
        markdown = "Texte **gras** et *italique* puis `code`.\n"
        doc = self.convert(markdown)

        paragraph = doc.paragraphs[0]
        self.assertEqual(paragraph.text, "Texte gras et italique puis code.")

        bold_run = next((run for run in paragraph.runs if run.text == "gras"), None)
        self.assertIsNotNone(bold_run)
        self.assertTrue(bold_run.bold)

        italic_run = next((run for run in paragraph.runs if run.text == "italique"), None)
        self.assertIsNotNone(italic_run)
        self.assertTrue(italic_run.italic)

        code_run = next((run for run in paragraph.runs if run.text == "code"), None)
        self.assertIsNotNone(code_run)
        self.assertEqual(code_run.font.name, "Courier New")

    def test_blockquote_lines_are_grouped_with_quote_indentation(self) -> None:
        markdown = "> Premiere ligne\n> Deuxieme ligne\n\nParagraphe normal\n"
        doc = self.convert(markdown)

        quote_paragraph = doc.paragraphs[0]
        self.assertEqual(quote_paragraph.text, "Premiere ligne Deuxieme ligne")

        left_indent = quote_paragraph.paragraph_format.left_indent
        self.assertIsNotNone(left_indent)
        self.assertGreater(left_indent, 700000)
        self.assertLess(left_indent, 730000)

        space_before = quote_paragraph.paragraph_format.space_before
        space_after = quote_paragraph.paragraph_format.space_after
        self.assertIsNotNone(space_before)
        self.assertIsNotNone(space_after)
        self.assertGreater(space_before, 110000)
        self.assertLess(space_before, 130000)
        self.assertGreater(space_after, 110000)
        self.assertLess(space_after, 130000)

        self.assertEqual(doc.paragraphs[1].text, "Paragraphe normal")

    def test_empty_blockquote_line_creates_new_quote_paragraph(self) -> None:
        markdown = "> Premier paragraphe\n>\n> Deuxieme paragraphe\n"
        doc = self.convert(markdown)

        self.assertEqual(doc.paragraphs[0].text, "Premier paragraphe")
        self.assertEqual(doc.paragraphs[1].text, "Deuxieme paragraphe")

        first_indent = doc.paragraphs[0].paragraph_format.left_indent
        second_indent = doc.paragraphs[1].paragraph_format.left_indent
        self.assertIsNotNone(first_indent)
        self.assertIsNotNone(second_indent)
        self.assertGreater(first_indent, 700000)
        self.assertGreater(second_indent, 700000)


if __name__ == "__main__":
    unittest.main()



