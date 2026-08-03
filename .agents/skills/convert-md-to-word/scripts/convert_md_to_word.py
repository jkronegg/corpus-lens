#!/usr/bin/env python3
"""Convertir un fichier Markdown en document Word (.docx)."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Cm


# Décale toutes les listes d'environ 1 cm vers la droite.
LIST_BASE_INDENT_CM = 1.0


def is_blockquote_line(line: str) -> bool:
    """Retourner True si la ligne Markdown appartient a un bloc citation."""
    return line.startswith('>')


def get_blockquote_content(line: str) -> str:
    """Extraire le contenu texte d'une ligne de citation (`>` optionnellement suivi d'un espace)."""
    content = line[1:]
    if content.startswith(' '):
        return content[1:]
    return content


def format_quote_paragraph(paragraph) -> None:
    """Appliquer la mise en forme de citation Word sur un paragraphe."""
    paragraph.paragraph_format.left_indent = 720000
    paragraph.paragraph_format.space_before = 120000
    paragraph.paragraph_format.space_after = 120000


def add_hyperlink(paragraph, text: str, url: str, force_italic: bool = False) -> None:
    """Ajouter un hyperlien cliquable à un paragraphe Word."""
    part = paragraph.part
    rel_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), rel_id)

    run = OxmlElement('w:r')
    run_props = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    run_props.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    run_props.append(underline)

    if force_italic:
        italic = OxmlElement('w:i')
        run_props.append(italic)

    run.append(run_props)

    text_node = OxmlElement('w:t')
    text_node.text = text
    run.append(text_node)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_formatted_text(paragraph, text: str, force_italic: bool = False) -> None:
    """Ajouter du texte formaté (gras, italique, liens, code) à un paragraphe."""
    pattern = r'(\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\)|`.*?`)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.italic = force_italic
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.italic = force_italic
        elif part.startswith('[') and '](' in part:
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                text_link = match.group(1)
                url = match.group(2)
                add_hyperlink(paragraph, text_link, url, force_italic=force_italic)
            else:
                run = paragraph.add_run(part)
                run.italic = force_italic
        else:
            run = paragraph.add_run(part)
            run.italic = force_italic


def parse_markdown_to_docx(md_file: str, docx_file: str) -> None:
    """Convertir un Markdown en document Word."""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    lines = content.split('\n')
    i = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph_buffer() -> None:
        """Écrire le paragraphe texte en attente (lignes Markdown fusionnées)."""
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        p = doc.add_paragraph()
        add_formatted_text(p, ' '.join(paragraph_buffer))
        paragraph_buffer = []

    while i < len(lines):
        line = lines[i]

        if line.startswith('# '):
            flush_paragraph_buffer()
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            flush_paragraph_buffer()
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            flush_paragraph_buffer()
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            flush_paragraph_buffer()
            doc.add_heading(line[5:].strip(), level=4)
        elif is_blockquote_line(line):
            flush_paragraph_buffer()
            quote_paragraph_lines: list[str] = []

            def flush_quote_paragraph() -> None:
                """Ecrire un paragraphe de citation en joignant ses lignes."""
                nonlocal quote_paragraph_lines
                if not quote_paragraph_lines:
                    return
                p = doc.add_paragraph()
                format_quote_paragraph(p)
                add_formatted_text(p, ' '.join(quote_paragraph_lines), force_italic=True)
                quote_paragraph_lines = []

            while i < len(lines) and is_blockquote_line(lines[i]):
                quote_line = get_blockquote_content(lines[i]).strip()
                if quote_line:
                    quote_paragraph_lines.append(quote_line)
                else:
                    # Une ligne `>` vide force un nouveau paragraphe dans la citation.
                    flush_quote_paragraph()
                i += 1
            i -= 1
            flush_quote_paragraph()
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            flush_paragraph_buffer()
            indent_level = (len(line) - len(line.lstrip())) // 2
            text = line.strip()[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(LIST_BASE_INDENT_CM + indent_level * 2.0)
            add_formatted_text(p, text)
        elif re.match(r'^\d+\)', line.strip()):
            flush_paragraph_buffer()
            indent_level = (len(line) - len(line.lstrip())) // 2
            text = re.sub(r'^\d+\)\s*', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(LIST_BASE_INDENT_CM + indent_level * 2.0)
            add_formatted_text(p, text)
        elif not line.strip():
            # En Markdown, une ligne vide sépare des paragraphes.
            flush_paragraph_buffer()
        else:
            paragraph_buffer.append(line.strip())

        i += 1

    flush_paragraph_buffer()

    doc.save(docx_file)
    print(f"Document créé : {docx_file}")


def infer_output_path(md_path: Path) -> Path:
    return md_path.with_suffix('.docx')


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Convertir un fichier Markdown en document Word (.docx)."
    )
    parser.add_argument(
        '--input',
        required=True,
        help='Chemin du fichier Markdown source.',
    )
    parser.add_argument(
        '--output',
        help='Chemin du fichier Word de sortie (.docx). Si omis, le même nom est utilisé.',
    )
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()

    md_path = Path(args.input).expanduser().resolve()
    if not md_path.exists():
        parser.error(f"Fichier Markdown introuvable: {md_path}")

    docx_path = Path(args.output).expanduser().resolve() if args.output else infer_output_path(md_path)
    parse_markdown_to_docx(str(md_path), str(docx_path))
    return 0


if __name__ == '__main__':
    raise SystemExit(main())

